import requests
import openai
import os
from docx import Document
import tkinter as tk
from tkinter import filedialog, messagebox

# 替换为你的 OpenAI API 密钥
openai.api_key = ''

# 创建窗口
root = tk.Tk()
root.title("网页抓取与 GPT 处理")
root.geometry("500x600")

# 标签和输入框
url_label = tk.Label(root, text="请输入网页 URL:")
url_label.pack(pady=10)
url_entry = tk.Entry(root, width=50)
url_entry.pack(pady=5)

# 输入自定义标题数量
title_count_label = tk.Label(root, text="请选择你要生成的标题数量 (1-5):")
title_count_label.pack(pady=10)
title_count_var = tk.IntVar(value=1)
title_count_entry = tk.Spinbox(root, from_=1, to=5, textvariable=title_count_var)
title_count_entry.pack(pady=5)

# 动态生成的标题输入框列表
title_widgets = []  # 存储标题标签和输入框的列表

# 动态生成标题输入框的函数
def generate_title_inputs():
    global title_widgets
    # 清除之前的标题标签和输入框
    for widget_set in title_widgets:
        for widget in widget_set:
            widget.destroy()
    title_widgets.clear()

    # 根据选择的标题数量生成输入框和标签
    title_count = title_count_var.get()
    for i in range(title_count):
        title_label = tk.Label(root, text=f"请输入第 {i+1} 个标题:")
        title_label.pack(pady=5)
        title_entry = tk.Entry(root, width=50)
        title_entry.pack(pady=5)
        title_widgets.append((title_label, title_entry))  # 将标签和输入框打包在一起保存

# 每次更改标题数量时重新生成输入框
title_count_entry.config(command=generate_title_inputs)

# 默认生成 1 个标题输入框
generate_title_inputs()

# 功能函数
def process_url():
    url = url_entry.get()  # 从输入框获取 URL
    title_count = title_count_var.get()  # 获取标题数量

    # 获取用户输入的标题
    titles = [entry.get() or f"标题 {i+1}" for i, (_, entry) in enumerate(title_widgets)]

    if not url:
        messagebox.showwarning("输入错误", "请输入有效的 URL")
        return

    try:
        # 请求网页内容
        response = requests.get(url)

        # 检查请求是否成功
        if response.status_code == 200:
            page_content = response.text

            # 根据标题动态生成 GPT 请求内容
            gpt_prompt = "請針對以下標題從網頁內容中提取相關信息，並以正式的語氣撰寫每個段落。請包括具體數據和事實，並保持語句簡潔流暢：\n\n"
            for i, title in enumerate(titles):
                gpt_prompt += f"{i+1}. {title}: 提取與此標題相關的重點資訊，並簡單列出要點。\n"
            gpt_prompt += f"\n原始網頁內容如下：\n\n{page_content}"

            # 调用 OpenAI 的 ChatCompletion 接口
            gpt_response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",  # 根据你的订阅选择 "gpt-3.5-turbo" 或 "gpt-4"
                messages=[
                    {"role": "system", "content": "你是一个乐于助人的助手。"},
                    {"role": "user", "content": gpt_prompt},  # 传递网页内容和结构提示
                ]
            )

            # 获取 GPT 的响应
            answer = gpt_response['choices'][0]['message']['content']

            # 选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word 文档", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加自定义的标题和内容
                content_sections = answer.split("\n\n")
                for i, title in enumerate(titles):
                    doc.add_heading(title, level=i+1)
                    if i < len(content_sections):
                        doc.add_paragraph(content_sections[i].strip())

                # 保存文档
                doc.save(save_path)
                messagebox.showinfo("成功", f"Word 文档已保存到: {save_path}")
        else:
            messagebox.showerror("错误", f"请求失败，状态码：{response.status_code}")
    except Exception as e:
        messagebox.showerror("错误", f"发生错误: {str(e)}")

# 按钮
process_button = tk.Button(root, text="抓取并处理网页", command=process_url)
process_button.pack(pady=20)

# 运行窗口
root.mainloop()
